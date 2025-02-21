import openpyxl
import argparse
import os
import time
import win32com.client  # Requires `pywin32`
from docx import Document
from docx.shared import Inches
import logging
from PIL import ImageGrab  # Requires `Pillow`

# Set max width (in inches) for images in Word
MAX_WIDTH_INCHES = 6.5  # Adjust as needed

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def export_excel_to_word(excel_file, sheet_name, word_file, section_title):
    # Check if files exist
    if not os.path.exists(excel_file):
        logger.error(f"Excel file '{excel_file}' not found.")
        return
    if not os.path.exists(word_file):
        logger.error(f"Word file '{word_file}' not found.")
        return

    try:
        # Open the Excel workbook
        wb = openpyxl.load_workbook(excel_file, data_only=True)
        if sheet_name not in wb.sheetnames:
            logger.error(f"Sheet '{sheet_name}' not found in the Excel file.")
            return
        sheet = wb[sheet_name]

        # Extract data from the sheet and format it as text (skip rows 1 to 4)
        data = []
        for row_index, row in enumerate(sheet.iter_rows(values_only=True)):
            if row_index < 4:  # Skip first 4 rows (index 0, 1, 2, 3)
                continue
            row_data = [str(cell) if cell is not None else "" for cell in row]
            if any(row_data):  # Skip row if all cells are empty
                data.append(row_data)

    except Exception as e:
        logger.error(f"Failed to open Excel file. {e}")
        return
    finally:
        wb.close()  # Close Excel to avoid lock issues

    try:
        # Load the Word document
        doc = Document(word_file)

        # Step 1: Locate the custom styled header and create new paragraphs
        insert_index = None
        for i, para in enumerate(doc.paragraphs):
            if para.style.name == "Custom" and section_title in para.text:
                insert_index = i + 1
                break
        else:
            logger.error(f"Header '{section_title}' with style 'Custom' not found in the Word document.")
            return

        # Open Excel to count images and maintain their order
        excel = win32com.client.Dispatch("Excel.Application")
        excel.Visible = False  # Run in the background
        wb_xl = excel.Workbooks.Open(os.path.abspath(excel_file))
        sheet_xl = wb_xl.Sheets[sheet_name]

        shapes = sheet_xl.Shapes  # Get all shapes (includes pasted screenshots)
        image_list = []

        for i in range(1, shapes.Count + 1):
            shape = shapes.Item(i)
            try:
                image_list.append((shape.TopLeftCell.Row, shape))  # Store row position
            except Exception as e:
                logger.warning(f"Could not determine row position for image '{shape.Name}'. Error: {e}")

        # Sort images by row position to maintain order
        image_list.sort(key=lambda x: x[0])
        total_items = len(data) + len(image_list)

        # Create new paragraphs for all items at the correct position
        for _ in range(total_items):
            doc.paragraphs[insert_index].insert_paragraph_before("")

        # Step 2: Locate the custom styled header again and insert data
        insert_index = None
        for i, para in enumerate(doc.paragraphs):
            if para.style.name == "Custom" and section_title in para.text:
                insert_index = i + 1
                break

        # Insert the first text paragraph
        if data:
            for cell in data[0]:
                if cell.strip():
                    doc.paragraphs[insert_index].add_run(cell)
                    insert_index += 1
            data_index = 1  # Start from the second text entry
        else:
            data_index = 0

        # Alternate inserting text and images
        image_index = 0
        while data_index < len(data) or image_index < len(image_list):
            if data_index < len(data):
                for cell in data[data_index]:
                    if cell.strip():
                        doc.paragraphs[insert_index].add_run(cell)
                        insert_index += 1
                data_index += 1
            
            if image_index < len(image_list):
                _, shape = image_list[image_index]
                image_path = f"temp_image_{shape.Name}.png"
                try:
                    shape.Copy()
                    time.sleep(0.5)  # Allow clipboard processing
                    image = ImageGrab.grabclipboard()
                    if image:
                        image.save(image_path, 'PNG')
                        doc.paragraphs[insert_index].add_run().add_picture(image_path, width=Inches(MAX_WIDTH_INCHES))
                        insert_index += 1
                        os.remove(image_path)  # Clean up temp image
                        logger.info(f"Image '{shape.Name}' added successfully.")
                    else:
                        logger.warning(f"Clipboard does not contain an image for '{shape.Name}'.")
                except Exception as e:
                    logger.warning(f"Could not process image '{shape.Name}'. Error: {e}")
                image_index += 1

        wb_xl.Close(False)
        excel.Quit()

    except Exception as e:
        logger.warning(f"Unable to extract pasted images (screenshots). Error: {e}")

    # Save the modified Word document
    try:
        doc.save(word_file)
        logger.info(f"Success: Data and images from '{sheet_name}' inserted under section '{section_title}' in '{word_file}'.")
    except Exception as e:
        logger.error(f"Failed to save the Word document. {e}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Export an Excel sheet to a specific section in a Word document, including images.")
    parser.add_argument("excel_file", help="Path to the Excel file")
    parser.add_argument("sheet_name", help="Name of the Excel sheet to export")
    parser.add_argument("word_file", help="Path to the Word document")
    parser.add_argument("section_title", help="Title of the section in the Word document where data should be inserted")

    args = parser.parse_args()
    export_excel_to_word(args.excel_file, args.sheet_name, args.word_file, args.section_title)
